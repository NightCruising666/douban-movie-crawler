#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
============================================================
生成实验报告 Word 文档
============================================================
读取清洗日志、可视化图片，生成格式化的 .docx 实验报告。
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import datetime


# ========== 用户信息（请修改为自己的信息）==========
STUDENT_NAME   = "待填写"
STUDENT_ID     = "待填写"
MAJOR          = "待填写"
CLASS_NAME     = "待填写"
# ==================================================


def set_cell_shading(cell, color_hex: str):
    """为表格单元格设置背景色。"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_heading_styled(doc, text: str, level: int = 1):
    """添加带格式的标题。"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    return heading


def add_code_block(doc, code_text: str):
    """以灰色背景段落模拟代码块。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    # 灰色底纹
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F0F0F0')
    shading.set(qn('w:val'), 'clear')
    p._element.get_or_add_pPr().append(shading)
    return p


def main():
    doc = Document()

    # ---- 页面设置 ----
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # ---- 样式设置 ----
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ================================================================
    # 封面
    # ================================================================
    for _ in range(6):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run('认知实习实验报告')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_p.add_run('—— 销售数据清洗与可视化分析')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    doc.add_paragraph()

    # 个人信息表格
    info_table = doc.add_table(rows=5, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ('专  业', MAJOR),
        ('班  级', CLASS_NAME),
        ('学  号', STUDENT_ID),
        ('姓  名', STUDENT_NAME),
        ('日  期', datetime.date.today().strftime('%Y年%m月%d日')),
    ]
    for i, (label, value) in enumerate(info_data):
        info_table.cell(i, 0).text = label
        info_table.cell(i, 1).text = value
        for j in range(2):
            cell = info_table.cell(i, j)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(14)
            if j == 0:
                set_cell_shading(cell, 'E8F0FE')
    info_table.style = 'Table Grid'

    doc.add_page_break()

    # ================================================================
    # 目录页（手动）
    # ================================================================
    add_heading_styled(doc, '目  录', level=1)
    toc_items = [
        '一、实验目的',
        '二、实验环境',
        '三、数据说明',
        '四、数据清洗过程',
        '    4.1 读取数据与EDA',
        '    4.2 处理 ERROR / UNKNOWN 值',
        '    4.3 处理缺失数据',
        '    4.4 剔除重复数据',
        '    4.5 数值列逻辑修复',
        '    4.6 保存清洗结果',
        '五、数据可视化分析',
        '六、实验总结',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(12)

    doc.add_page_break()

    # ================================================================
    # 一、实验目的
    # ================================================================
    add_heading_styled(doc, '一、实验目的', level=1)

    objectives = [
        '掌握使用 Python (Pandas) 读取和清洗 CSV 数据的基本方法；',
        '学会识别和处理缺失值、异常值（ERROR/UNKNOWN）以及重复数据；',
        '掌握数据可视化工具（Matplotlib / Seaborn）的使用，能够对不同类别产品的价格、购买数量等统计特征进行图形化展示；',
        '培养数据敏感性，能够对清洗前后的数据做出合理的对比分析和解释；',
        '完成从原始数据到干净数据集、再到可视化报告的全流程实践。',
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')

    # ================================================================
    # 二、实验环境
    # ================================================================
    add_heading_styled(doc, '二、实验环境', level=1)

    env_data = [
        ('操作系统', 'macOS / Windows'),
        ('编程语言', 'Python 3.x'),
        ('主要库', 'pandas, numpy, matplotlib, seaborn, python-docx'),
        ('开发工具', 'VS Code / Jupyter Notebook'),
        ('数据文件', 'sales.csv（原始）→ cleaned_sales.csv（清洗后）'),
    ]
    env_table = doc.add_table(rows=len(env_data), cols=2)
    env_table.style = 'Table Grid'
    env_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(env_data):
        env_table.cell(i, 0).text = k
        env_table.cell(i, 1).text = v
        set_cell_shading(env_table.cell(i, 0), 'E8F0FE')
        for j in range(2):
            for p in env_table.cell(i, j).paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()

    # ================================================================
    # 三、数据说明
    # ================================================================
    add_heading_styled(doc, '三、数据说明', level=1)

    doc.add_paragraph(
        '本实验使用的原始数据文件为 sales.csv，包含 10,000 条交易记录，共 8 个字段：'
    )

    col_data = [
        ('Transaction ID', '交易唯一标识符', 'TXN_xxxxxxx'),
        ('Item', '产品名称', 'Coffee, Cake, Cookie, Salad, Smoothie, Sandwich, Juice, Tea'),
        ('Quantity', '购买数量', '1 ~ 5'),
        ('Price Per Unit', '单价', '1.0, 1.5, 2.0, 3.0, 4.0, 5.0'),
        ('Total Spent', '总消费金额', 'Quantity × Price Per Unit'),
        ('Payment Method', '支付方式', 'Credit Card, Cash, Digital Wallet'),
        ('Location', '交易地点', 'Takeaway（外卖）, In-store（堂食）'),
        ('Transaction Date', '交易日期', '2023-01-01 ~ 2023-12-31'),
    ]
    col_table = doc.add_table(rows=len(col_data) + 1, cols=3)
    col_table.style = 'Table Grid'
    col_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['字段名', '含义', '正常取值范围']
    for j, h in enumerate(headers):
        col_table.cell(0, j).text = h
        set_cell_shading(col_table.cell(0, j), '1A3C6E')
        for p in col_table.cell(0, j).paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(9)
    for i, (col, desc, val_range) in enumerate(col_data):
        col_table.cell(i + 1, 0).text = col
        col_table.cell(i + 1, 1).text = desc
        col_table.cell(i + 1, 2).text = val_range
        for j in range(3):
            for p in col_table.cell(i + 1, j).paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()

    doc.add_paragraph(
        '数据中存在以下问题需要处理：'
        '(1) 部分单元格为空值（NaN 或空字符串）；'
        '(2) 部分单元格的值为 "ERROR" 或 "UNKNOWN"，属于无效数据；'
        '(3) 数值列中存在文本型错误值，导致无法直接进行计算；'
        '(4) 部分行的 Transaction ID 可能重复。'
    )

    # ================================================================
    # 四、数据清洗过程
    # ================================================================
    add_heading_styled(doc, '四、数据清洗过程', level=1)

    # --- 4.1 ---
    add_heading_styled(doc, '4.1 读取数据与探索性数据分析（EDA）', level=2)

    doc.add_paragraph(
        '首先使用 pandas.read_csv() 以字符串类型读取全部数据，避免自动类型转换导致的信息丢失。'
        '读取后共获得 10,000 行 × 8 列的数据。随后进行以下 EDA 检查：'
    )

    doc.add_paragraph('各列缺失值（NaN）统计：', style='List Bullet')
    doc.add_paragraph('    Item: 333 个 | Quantity: 138 个 | Price Per Unit: 179 个 | '
                      'Total Spent: 173 个 | Payment Method: 2,579 个 | '
                      'Location: 3,265 个 | Transaction Date: 159 个')

    doc.add_paragraph('各列 "ERROR" 值统计：', style='List Bullet')
    doc.add_paragraph('    Item: 292 个 | Quantity: 170 个 | Price Per Unit: 190 个 | '
                      'Total Spent: 164 个 | Payment Method: 306 个 | '
                      'Location: 358 个 | Transaction Date: 142 个')

    doc.add_paragraph('各列 "UNKNOWN" 值统计：', style='List Bullet')
    doc.add_paragraph('    Item: 344 个 | Quantity: 171 个 | Price Per Unit: 164 个 | '
                      'Total Spent: 165 个 | Payment Method: 293 个 | '
                      'Location: 338 个 | Transaction Date: 159 个')

    doc.add_paragraph('重复行检查：', style='List Bullet')
    doc.add_paragraph('    完全重复行数: 0 | Transaction ID 重复数: 0')

    # --- 4.2 ---
    add_heading_styled(doc, '4.2 处理 ERROR / UNKNOWN 无效数据', level=2)

    doc.add_paragraph(
        '遍历每一列，将所有值为 "ERROR" 或 "UNKNOWN"（不区分大小写）的单元格替换为 NumPy NaN。'
        '这一步将文本形式的异常值统一转换为标准的缺失值表示，方便后续统一处理。'
    )
    add_code_block(doc,
        '# 替换 ERROR / UNKNOWN 为 NaN\n'
        'for col in df.columns:\n'
        '    mask = df[col].str.upper().isin([\'ERROR\', \'UNKNOWN\'])\n'
        '    df.loc[mask, col] = np.nan'
    )
    doc.add_paragraph(
        '例如：第4行 Total Spent = "ERROR" → NaN；第5行 Payment Method & Location = "UNKNOWN" → NaN；'
        '第33行 Item = "UNKNOWN", Total Spent = "ERROR", Location = "ERROR" → 全部替换为 NaN。'
    )

    # --- 4.3 ---
    add_heading_styled(doc, '4.3 处理缺失数据', level=2)

    doc.add_paragraph('缺失数据的处理策略分为以下几步：')

    doc.add_paragraph('空字符串 → NaN：将各列中的空字符串统一替换为 NaN。', style='List Bullet')

    doc.add_paragraph('数值列格式转换：将 Quantity、Price Per Unit、Total Spent 转换为数值类型，'
                      '无法转换的（如残留的异常文本）自动变为 NaN。', style='List Bullet')

    doc.add_paragraph('日期列格式转换：将 Transaction Date 转换为 datetime 类型，'
                      '无效日期变为 NaT。', style='List Bullet')

    doc.add_paragraph('数值列逻辑修复（利用三列之间的算术关系）：', style='List Bullet')
    doc.add_paragraph(
        '    • 若 Total Spent 缺失但 Quantity 和 Price Per Unit 完整 → '
        'Total Spent = Quantity × Price Per Unit；\n'
        '    • 若 Quantity 缺失但 Total Spent 和 Price Per Unit 完整 → '
        'Quantity = Total Spent ÷ Price Per Unit；\n'
        '    • 若 Price Per Unit 缺失但 Total Spent 和 Quantity 完整 → '
        'Price Per Unit = Total Spent ÷ Quantity。'
    )

    doc.add_paragraph('剔除仍含缺失值的行：对修复后仍存在 NaN 的行进行剔除（dropna）。', style='List Bullet')
    doc.add_paragraph(
        '残留 NaN 的主要来源是：Item / Payment Method / Location 等类别列无法通过逻辑修复补全，'
        '且这些字段对分析至关重要，因此选择删除。'
    )

    # --- 4.4 ---
    add_heading_styled(doc, '4.4 剔除重复数据', level=2)

    doc.add_paragraph(
        '经检查，原始数据中无完全重复的行，也无重复的 Transaction ID。'
        '但为保障数据质量，代码中仍包含 drop_duplicates() 步骤作为防御性编程。'
    )

    # --- 4.5 ---
    add_heading_styled(doc, '4.5 数值列逻辑修复', level=2)

    doc.add_paragraph(
        '由于 Quantity、Price Per Unit、Total Spent 三列之间存在明确的数学关系 '
        '(Total Spent = Quantity × Price Per Unit)，可以利用这一约束对部分缺失值进行补全。'
    )
    add_code_block(doc,
        '# 补全 Total Spent\n'
        'mask = df[\'Total Spent\'].isnull() & df[\'Quantity\'].notnull() & df[\'Price Per Unit\'].notnull()\n'
        'df.loc[mask, \'Total Spent\'] = df.loc[mask, \'Quantity\'] * df.loc[mask, \'Price Per Unit\']\n'
        '\n'
        '# 补全 Quantity\n'
        'mask = df[\'Quantity\'].isnull() & df[\'Total Spent\'].notnull() & df[\'Price Per Unit\'].notnull()\n'
        'df.loc[mask, \'Quantity\'] = df.loc[mask, \'Total Spent\'] / df.loc[mask, \'Price Per Unit\']'
    )

    # --- 4.6 ---
    add_heading_styled(doc, '4.6 保存清洗结果', level=2)

    doc.add_paragraph(
        '清洗完成后，将数据保存为 cleaned_sales.csv（UTF-8 with BOM 编码，确保在 Excel 中正确显示中文）。'
        '最终数据类型为：Quantity (int)、Price Per Unit (float)、Total Spent (float)、'
        'Transaction Date (datetime)，其余为字符串类型。'
    )

    doc.add_paragraph(
        '清洗效果：从原始 10,000 行数据中移除了无效数据行，最终保留有效记录。'
        '具体保留行数请参见程序运行日志。'
    )

    # ================================================================
    # 五、数据可视化分析
    # ================================================================
    add_heading_styled(doc, '五、数据可视化分析', level=1)

    # --- 图1 ---
    add_heading_styled(doc, '5.1 各类产品销售统计', level=2)
    doc.add_paragraph(
        '图1 展示了各类产品的销售总量、交易笔数和平均单价。'
        '通过柱状图可以直观对比不同产品的市场表现：'
        '销量最高的产品反映了消费者的主要偏好，'
        '而平均单价的差异则体现了产品的定位层级。'
    )
    img_path = 'fig1_product_sales.png'
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('图1  各类产品销售总量、交易笔数及平均单价对比')

    # --- 图2 ---
    add_heading_styled(doc, '5.2 价格与购买数量分布', level=2)
    doc.add_paragraph(
        '图2 使用箱线图展示了各类产品单价和购买数量的分布情况。'
        '箱线图能够直观地显示数据的中位数、四分位数范围以及异常值，'
        '有助于识别不同产品在价格和销量上的离散程度。'
    )
    img_path = 'fig2_boxplot_price_qty.png'
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('图2  各类产品单价与购买数量箱线图')

    # --- 图3 ---
    add_heading_styled(doc, '5.3 销售额占比与支付方式', level=2)
    doc.add_paragraph(
        '图3 左侧饼图展示了各类产品在总销售额中的占比，可以清晰看出哪些产品是营收主力。'
        '右侧柱状图展示了支付方式的分布，反映了消费者的支付习惯。'
    )
    img_path = 'fig3_revenue_payment.png'
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('图3  各类产品销售额占比与支付方式分布')

    # --- 图4 ---
    add_heading_styled(doc, '5.4 月度销售趋势与位置分布', level=2)
    doc.add_paragraph(
        '图4 左侧折线图展示了 2023 年各月的销售额变化趋势，'
        '有助于识别销售旺季与淡季。'
        '右侧柱状图展示了交易位置的分布（堂食 vs 外卖）。'
    )
    img_path = 'fig4_trend_location.png'
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('图4  月度销售额趋势与交易位置分布')

    # --- 图5 ---
    add_heading_styled(doc, '5.5 交叉统计热力图', level=2)
    doc.add_paragraph(
        '图5 使用热力图展示了产品类别 × 支付方式、产品类别 × 位置的交叉统计。'
        '颜色深浅表示交易笔数的多少，便于发现不同维度之间的关联模式，'
        '例如某些产品是否更倾向于特定的支付方式或消费场景。'
    )
    img_path = 'fig5_heatmap_cross.png'
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('图5  产品类别与支付方式 / 位置的交叉统计热力图')

    # --- 图6 ---
    add_heading_styled(doc, '5.6 描述性统计汇总', level=2)
    doc.add_paragraph(
        '图6 以表格形式汇总了各类产品的描述性统计指标，'
        '包括交易笔数、销售总量、平均单价、销售额总和、平均每笔消费和平均每次购买量。'
        '该表为数据分析提供了全面的定量参考。'
    )
    img_path = 'fig6_summary_table.png'
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('图6  各类产品描述性统计汇总表')

    # ================================================================
    # 六、实验总结
    # ================================================================
    doc.add_page_break()
    add_heading_styled(doc, '六、实验总结', level=1)

    summary_points = [
        '数据清洗是数据分析的前提和基础。本实验通过对 sales.csv 的系统清洗，'
        '将含有大量 ERROR、UNKNOWN、空值等问题的原始数据转化为格式规范、内容完整的 cleaned_sales.csv。',

        '处理策略的选择需要结合实际数据特征：'
        '对于可逻辑推断的数值缺失（如 Total Spent = Quantity × Price Per Unit），'
        '优先采用计算补全；对于类别型缺失，因无法可靠推断而选择删除。'
        '这种"能修则修、不能修则删"的策略兼顾了数据完整性和准确性。',

        '可视化分析揭示了若干有价值的信息：'
        '不同产品的销量、单价和销售额差异显著，可以据此优化产品组合和定价策略；'
        '支付方式以某一种为主，反映了消费者的支付偏好；'
        '月度销售趋势显示出一定季节性波动，可为库存管理和促销活动提供参考。',

        '本实验完整实践了"原始数据 → 探索分析 → 清洗处理 → 可视化展示 → 报告撰写"的全流程，'
        '有效提升了数据分析和 Python 编程的实践能力。',
    ]
    for sp in summary_points:
        doc.add_paragraph(sp)

    doc.add_paragraph()
    doc.add_paragraph('=' * 50)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f'报告人：{STUDENT_NAME}\n日期：{datetime.date.today().strftime("%Y年%m月%d日")}')
    run.font.size = Pt(11)

    # ================================================================
    # 保存
    # ================================================================
    output_path = '实验报告_销售数据清洗与可视化.docx'
    doc.save(output_path)
    print(f'实验报告已生成: {output_path}')


if __name__ == '__main__':
    main()
